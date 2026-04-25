# ============================================
# RESUME ANALYZER — Main Application
# ============================================

import streamlit as st
import pdfplumber
from docx import Document
import spacy
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re

# ReportLab for PDF report generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from io import BytesIO

# ============================================
# PAGE CONFIGURATION (must be the FIRST Streamlit call)
# ============================================
st.set_page_config(
    page_title="Resume Analyzer",
    page_icon="📄",
    layout="wide"
)

# ============================================
# CUSTOM CSS FOR POLISHED UI
# ============================================
st.markdown("""
<style>
    h1 {
        font-weight: 700 !important;
        letter-spacing: -0.02em !important;
    }
    h2, h3 {
        font-weight: 600 !important;
        margin-top: 1.5rem !important;
    }
    [data-testid="stMetricValue"] {
        font-size: 3rem !important;
        font-weight: 700 !important;
    }
    .stButton > button[kind="primary"],
    .stDownloadButton > button[kind="primary"] {
        font-size: 1.1rem !important;
        font-weight: 600 !important;
        padding: 0.7rem !important;
        border-radius: 10px !important;
    }
    code {
        background-color: rgba(131, 165, 152, 0.15) !important;
        color: #a8c686 !important;
        padding: 2px 8px !important;
        border-radius: 6px !important;
        font-size: 0.9rem !important;
    }
    hr {
        margin: 2rem 0 !important;
        opacity: 0.3 !important;
    }
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .streamlit-expanderHeader {
        font-weight: 600 !important;
    }
</style>
""", unsafe_allow_html=True)

# ============================================
# Load spaCy model once (cached for performance)
# ============================================
@st.cache_resource
def load_nlp_model():
    """Load the spaCy English language model."""
    return spacy.load("en_core_web_sm")

nlp = load_nlp_model()

# ============================================
# SKILL DATABASE
# Covers both software/tech and electrical/electronics skills
# ============================================
SKILLS_DB = [
    # Programming languages
    "python", "java", "javascript", "typescript", "c++", "c#", "c", "go", "rust",
    "kotlin", "swift", "ruby", "php", "r", "scala", "matlab", "embedded c",
    "verilog", "vhdl", "assembly",
    # Web technologies
    "html", "css", "react", "angular", "vue", "node.js", "express", "next.js",
    "django", "flask", "fastapi", "spring", "rails",
    # Databases
    "sql", "mysql", "postgresql", "mongodb", "redis", "sqlite", "oracle",
    "cassandra", "dynamodb", "firebase",
    # Cloud & DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "terraform",
    "ansible", "ci/cd", "git", "github", "gitlab",
    # Data & ML
    "machine learning", "deep learning", "tensorflow", "pytorch", "keras",
    "scikit-learn", "pandas", "numpy", "nlp", "computer vision",
    "data analysis", "data visualization", "tableau", "power bi", "excel",
    "statistics", "a/b testing", "big data", "hadoop", "spark",
    "jupyter notebook", "pycharm", "vs code",
    # Other software tech
    "rest api", "graphql", "microservices", "agile", "scrum", "jira",
    "linux", "bash", "shell scripting", "figma",
    # Electrical & electronics engineering
    "circuit design", "pcb design", "pcb", "electrical design",
    "power electronics", "power systems", "power system protection",
    "electrical wiring", "wiring", "wiring systems", "electrical installation",
    "electrical fixtures", "electrical codes", "blueprints", "schematics",
    "technical drawings", "autocad", "autocad electrical", "ltspice",
    "eagle", "altium", "kicad", "proteus", "multisim",
    "simulink", "signal processing", "control systems", "control theory",
    "pid controller", "pid", "feedback control", "closed-loop control",
    "pwm", "pwm control", "inverters", "converters", "dc-dc converter",
    "ac-dc converter", "rectifiers", "transformers",
    "microcontroller", "microcontrollers", "arduino", "arduino uno",
    "raspberry pi", "esp32", "stm32", "fpga", "plc", "scada",
    "embedded systems", "iot", "internet of things", "sensors",
    "ultrasonic sensor", "zero-crossing detection",
    "ac machines", "induction motor", "dc motor", "stepper motor",
    "servo motor", "relay", "actuator", "igbt", "mosfet", "bjt",
    "oscilloscope", "multimeter", "function generator", "logic analyzer",
    "soldering", "troubleshooting", "debugging", "repair",
    "electrical safety", "safety compliance", "installation", "maintenance",
    # Electrical domain knowledge
    "kirchhoff", "ohms law", "fourier", "laplace", "frequency response",
    "digital signal processing", "dsp", "analog electronics", "digital electronics",
    # Soft skills
    "communication", "teamwork", "leadership", "problem solving",
    "analytical thinking", "project management", "collaboration",
    "attention to detail", "time management",
]

# ============================================
# FILE EXTRACTION FUNCTIONS
# ============================================

def extract_text_from_pdf(pdf_file):
    """
    Extract text from a PDF using pdfplumber.
    pdfplumber is better than PyPDF2 at preserving word spacing,
    especially for PDFs with complex fonts (LaTeX, designer templates).
    """
    try:
        text = ""
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        st.error(f"Could not read the PDF: {e}")
        return ""


def extract_text_from_docx(docx_file):
    """
    Extract text from a DOCX (Word) file using python-docx.
    Reads paragraph text and basic table contents.
    """
    try:
        text = ""
        doc = Document(docx_file)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        return text
    except Exception as e:
        st.error(f"Could not read the DOCX: {e}")
        return ""


def extract_text_from_file(uploaded_file):
    """Router: detect file type and call the right extractor."""
    if uploaded_file is None:
        return ""
    filename = uploaded_file.name.lower()
    if filename.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    else:
        st.error(f"Unsupported file type: {filename}. Please upload a PDF or DOCX.")
        return ""


# ============================================
# ANALYSIS HELPER FUNCTIONS
# ============================================

def clean_text(text):
    """Lowercase the text and remove extra whitespace and weird characters."""
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^a-z0-9+.#/\s-]', ' ', text)
    return text.strip()


def extract_skills(text, skills_db):
    """
    Find which skills from our database appear in the text.
    Uses both normal and 'compact' matching to handle PDF extraction quirks.
    """
    text_clean = clean_text(text)
    text_compact = re.sub(r'\s+', '', text_clean)
    found = set()
    for skill in skills_db:
        skill_clean = skill.lower()
        if (" " in skill_clean or "." in skill_clean or "+" in skill_clean
                or "#" in skill_clean or "/" in skill_clean or "-" in skill_clean):
            if skill_clean in text_clean:
                found.add(skill)
            elif len(skill_clean) >= 5 and skill_clean.replace(" ", "") in text_compact:
                found.add(skill)
        else:
            pattern = r'\b' + re.escape(skill_clean) + r'\b'
            if re.search(pattern, text_clean):
                found.add(skill)
            elif len(skill_clean) >= 5 and skill_clean in text_compact:
                found.add(skill)
    return found


def calculate_similarity(resume_text, jd_text):
    """Use TF-IDF and cosine similarity to compute a match score (0–100)."""
    documents = [clean_text(resume_text), clean_text(jd_text)]
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(documents)
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    return round(similarity * 100, 2)


def get_improvement_tips(missing_skills, score):
    """Generate a friendly list of suggestions, calibrated to realistic TF-IDF scores."""
    tips = []
    if score < 15:
        tips.append("📉 Your match score is low. The resume and JD seem to cover quite different domains.")
    elif score < 30:
        tips.append("📊 Moderate match — there's room to better align your resume with this job description.")
    elif score < 50:
        tips.append("✅ Strong match! Your resume aligns well with this role.")
    else:
        tips.append("🌟 Excellent match — your resume is very well-tailored for this role.")

    if missing_skills:
        top_missing = list(missing_skills)[:8]
        tips.append(
            f"🎯 Consider adding these skills if you genuinely have them: **{', '.join(top_missing)}**"
        )
        tips.append(
            "💡 Tip: Don't just list skills — back them up with a project or bullet point showing you've used them."
        )

    tips.append("📝 Use keywords *verbatim* from the job description when possible. Many companies filter resumes using ATS (Applicant Tracking Systems) that do exact keyword matching.")
    return tips


# ============================================
# REPORT GENERATION
# ============================================

def generate_report(score, matched_skills, missing_skills, matched_keywords,
                    missing_keywords, tips, resume_filename="resume"):
    """Build a nicely-formatted plain-text report."""
    lines = []
    lines.append("=" * 60)
    lines.append("RESUME ANALYSIS REPORT")
    lines.append("=" * 60)
    lines.append(f"\nResume analyzed: {resume_filename}")
    lines.append(f"Match Score: {score}%\n")

    lines.append("-" * 60)
    lines.append("MATCHING SKILLS")
    lines.append("-" * 60)
    if matched_skills:
        for s in sorted(matched_skills):
            lines.append(f"  + {s}")
    else:
        lines.append("  (None detected from skills database)")

    lines.append("")
    lines.append("-" * 60)
    lines.append("MISSING SKILLS (in JD but not in resume)")
    lines.append("-" * 60)
    if missing_skills:
        for s in sorted(missing_skills):
            lines.append(f"  - {s}")
    else:
        lines.append("  (No missing skills from our database)")

    lines.append("")
    lines.append("-" * 60)
    lines.append("TOP KEYWORDS FROM JOB DESCRIPTION")
    lines.append("-" * 60)
    lines.append(f"\nFound in your resume:")
    lines.append("  " + (", ".join(matched_keywords) if matched_keywords else "(none)"))
    lines.append(f"\nMissing from your resume:")
    lines.append("  " + (", ".join(missing_keywords) if missing_keywords else "(none)"))

    lines.append("")
    lines.append("-" * 60)
    lines.append("SUGGESTIONS")
    lines.append("-" * 60)
    for tip in tips:
        clean_tip = tip.replace("**", "").replace("*", "")
        lines.append(f"  - {clean_tip}")

    lines.append("")
    lines.append("=" * 60)
    lines.append("Generated by Resume Analyzer")
    lines.append("=" * 60)
    return "\n".join(lines)


def generate_pdf_report(score, matched_skills, missing_skills, matched_keywords,
                        missing_keywords, tips, resume_filename="resume"):
    """
    Build a styled PDF report using ReportLab.
    Strips emojis from text since default PDF fonts don't support them.
    Returns the PDF as bytes.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.6 * inch,
        leftMargin=0.6 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch
    )

    # Helper to clean text for PDF (emoji strip + markdown conversion)
    def clean_for_pdf(text):
        """Remove emojis but keep useful Unicode (em-dash, en-dash, smart quotes)."""
        # Replace useful Unicode chars with ASCII equivalents BEFORE stripping
        replacements = {
            '\u2014': '--',   # em-dash —
            '\u2013': '-',    # en-dash –
            '\u2018': "'",    # left single quote '
            '\u2019': "'",    # right single quote '
            '\u201c': '"',    # left double quote "
            '\u201d': '"',    # right double quote "
            '\u2026': '...',  # ellipsis …
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        # Now strip remaining non-ASCII (emojis, etc.)
        text = re.sub(r'[^\x00-\x7F]+', '', text)
        # Convert markdown
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
        return text.strip()

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=22,
        textColor=HexColor('#1e3a5f'),
        spaceAfter=6,
        alignment=0,
    )
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=HexColor('#666666'),
        spaceAfter=16,
    )
    section_heading = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=HexColor('#2d5a7b'),
        spaceBefore=14,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=4,
    )
    score_style = ParagraphStyle(
        'Score',
        parent=styles['Heading1'],
        fontSize=40,
        alignment=0,
        spaceAfter=0,   # no auto-space after; we control spacing manually
        leading=44,     # line height (must be >= fontSize for big text)
    )

    # Color-code the score (calibrated to realistic TF-IDF ranges)
    if score < 15:
        score_color = HexColor('#c0392b')  # red
        score_label = "Poor Match"
    elif score < 30:
        score_color = HexColor('#d68910')  # amber
        score_label = "Moderate Match"
    elif score < 50:
        score_color = HexColor('#1e8449')  # green
        score_label = "Strong Match"
    else:
        score_color = HexColor('#0e6655')  # darker green
        score_label = "Excellent Match"
    score_style.textColor = score_color

    label_style = ParagraphStyle(
        'ScoreLabel',
        parent=styles['Normal'],
        fontSize=11,
        textColor=score_color,
        spaceAfter=10,
    )

    story = []
    story.append(Paragraph("Resume Analysis Report", title_style))
    story.append(Paragraph(f"Resume: <b>{resume_filename}</b>", subtitle_style))

    # Match Score
    story.append(Paragraph("Match Score", section_heading))
    story.append(Paragraph(f"{score}%", score_style))
    story.append(Spacer(1, 0.05 * inch))  # small gap between score and label
    story.append(Paragraph(f"<b>{score_label}</b>", label_style))
    story.append(Spacer(1, 0.15 * inch))  # bigger gap before next section

    # Matching Skills
    story.append(Paragraph("Matching Skills", section_heading))
    if matched_skills:
        story.append(Paragraph(", ".join(sorted(matched_skills)), body_style))
    else:
        story.append(Paragraph("<i>None detected from skills database.</i>", body_style))

    # Missing Skills
    story.append(Paragraph("Missing Skills (in JD but not in resume)", section_heading))
    if missing_skills:
        story.append(Paragraph(", ".join(sorted(missing_skills)), body_style))
    else:
        story.append(Paragraph("<i>No missing skills from our database.</i>", body_style))

    # Top Keywords
    story.append(Paragraph("Top Keywords from Job Description", section_heading))
    story.append(Paragraph("<b>Found in your resume:</b>", body_style))
    if matched_keywords:
        story.append(Paragraph(", ".join(matched_keywords), body_style))
    else:
        story.append(Paragraph("<i>None.</i>", body_style))
    story.append(Spacer(1, 0.08 * inch))
    story.append(Paragraph("<b>Missing from your resume:</b>", body_style))
    if missing_keywords:
        story.append(Paragraph(", ".join(missing_keywords), body_style))
    else:
        story.append(Paragraph("<i>None.</i>", body_style))

    # Suggestions
    story.append(Paragraph("Suggestions to Improve", section_heading))
    for tip in tips:
        clean_tip = clean_for_pdf(tip)
        story.append(Paragraph(f"&bull; {clean_tip}", body_style))

    # Footer
    story.append(Spacer(1, 0.3 * inch))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=HexColor('#999999'),
        alignment=1,
    )
    story.append(Paragraph("Generated by Resume Analyzer", footer_style))

    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes


# ============================================
# STREAMLIT UI
# ============================================

# Header banner
st.markdown("""
<div style="
    background: linear-gradient(135deg, #1e3a5f 0%, #2d5a7b 100%);
    padding: 2rem;
    border-radius: 12px;
    margin-bottom: 2rem;
">
    <h1 style="color: white; margin: 0; font-size: 2.5rem;">
        📄 Resume Analyzer
    </h1>
    <p style="color: rgba(255,255,255,0.85); margin: 0.5rem 0 0 0; font-size: 1.1rem;">
        AI-powered matching between your resume and job descriptions.
        Get a score, find missing skills, and get improvement tips instantly.
    </p>
</div>
""", unsafe_allow_html=True)

# Two-column layout for inputs
col1, col2 = st.columns(2)

with col1:
    st.subheader("📎 Step 1: Upload Your Resume")
    uploaded_file = st.file_uploader(
        "Choose a PDF or DOCX file",
        type=["pdf", "docx"],
        help="PDF and Word documents are supported."
    )

with col2:
    st.subheader("📋 Step 2: Paste the Job Description")
    job_description = st.text_area(
        "Paste the full job description here",
        height=250,
        placeholder="E.g., We are looking for a Python developer with experience in..."
    )

st.divider()
analyze_clicked = st.button("🔍 Analyze", type="primary", use_container_width=True)

# ============================================
# RESULTS SECTION
# ============================================

if analyze_clicked:
    if not uploaded_file:
        st.warning("⚠️ Please upload a resume (PDF or DOCX) first.")
    elif not job_description.strip():
        st.warning("⚠️ Please paste a job description.")
    else:
        with st.spinner("Analyzing... this takes a few seconds."):
            resume_text = extract_text_from_file(uploaded_file)

            if not resume_text.strip():
                st.error("Couldn't extract any text. The file might be a scanned image or empty.")
            else:
                score = calculate_similarity(resume_text, job_description)
                resume_skills = extract_skills(resume_text, SKILLS_DB)
                jd_skills = extract_skills(job_description, SKILLS_DB)
                matched_skills = resume_skills & jd_skills
                missing_skills = jd_skills - resume_skills

                # Match score
                st.subheader("📊 Match Score")
                st.metric(label="Overall Match", value=f"{score}%")
                st.progress(score / 100)

                # Skills breakdown
                skill_col1, skill_col2 = st.columns(2)
                with skill_col1:
                    st.subheader("✅ Matching Skills")
                    if matched_skills:
                        for skill in sorted(matched_skills):
                            st.markdown(f"- {skill}")
                    else:
                        st.info("No matching skills found from our database.")

                with skill_col2:
                    st.subheader("❌ Missing Skills")
                    if missing_skills:
                        for skill in sorted(missing_skills):
                            st.markdown(f"- {skill}")
                    elif jd_skills:
                        st.success("You've covered all skills from the JD that are in our database!")
                    else:
                        st.info(
                            "No skills from our database were detected in the job description. "
                            "See the general keyword analysis below."
                        )

                # Top keywords analysis
                st.subheader("🔑 Top Keywords from Job Description")
                st.caption(
                    "These are the most important words in the JD (based on frequency & uniqueness). "
                    "Missing ones could be worth adding to your resume."
                )

                matched_keywords = []
                missing_keywords = []
                try:
                    keyword_vectorizer = TfidfVectorizer(
                        stop_words='english',
                        max_features=20,
                        ngram_range=(1, 2)
                    )
                    keyword_vectorizer.fit_transform([clean_text(job_description)])
                    top_keywords = keyword_vectorizer.get_feature_names_out()
                    resume_text_clean = clean_text(resume_text)
                    resume_text_compact = re.sub(r'\s+', '', resume_text_clean)
                    for kw in top_keywords:
                        kw_compact = kw.replace(" ", "")
                        if kw in resume_text_clean or (len(kw_compact) >= 5 and kw_compact in resume_text_compact):
                            matched_keywords.append(kw)
                        else:
                            missing_keywords.append(kw)

                    kw_col1, kw_col2 = st.columns(2)
                    with kw_col1:
                        st.markdown("**✅ Found in your resume:**")
                        if matched_keywords:
                            st.markdown(", ".join(f"`{k}`" for k in matched_keywords))
                        else:
                            st.caption("None of the top JD keywords appear in your resume.")
                    with kw_col2:
                        st.markdown("**⚠️ Missing from your resume:**")
                        if missing_keywords:
                            st.markdown(", ".join(f"`{k}`" for k in missing_keywords))
                        else:
                            st.caption("You've included all top keywords.")
                except Exception as e:
                    st.caption(f"Could not run keyword analysis: {e}")

                # Suggestions
                st.subheader("💡 Suggestions to Improve")
                tips = get_improvement_tips(missing_skills, score)
                for tip in tips:
                    st.markdown(tip)

                # Download report section
                st.divider()
                st.subheader("📥 Download Your Report")

                report_text = generate_report(
                    score=score,
                    matched_skills=matched_skills,
                    missing_skills=missing_skills,
                    matched_keywords=matched_keywords,
                    missing_keywords=missing_keywords,
                    tips=tips,
                    resume_filename=uploaded_file.name
                )

                pdf_bytes = generate_pdf_report(
                    score=score,
                    matched_skills=matched_skills,
                    missing_skills=missing_skills,
                    matched_keywords=matched_keywords,
                    missing_keywords=missing_keywords,
                    tips=tips,
                    resume_filename=uploaded_file.name
                )

                download_col1, download_col2 = st.columns(2)
                with download_col1:
                    st.download_button(
                        label="📄 Download as PDF",
                        data=pdf_bytes,
                        file_name="resume_analysis_report.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        type="primary"
                    )
                with download_col2:
                    st.download_button(
                        label="📝 Download as TXT",
                        data=report_text,
                        file_name="resume_analysis_report.txt",
                        mime="text/plain",
                        use_container_width=True
                    )

                # Debug section
                with st.expander("🔧 Debug info — what the analyzer detected"):
                    debug_col1, debug_col2 = st.columns(2)
                    with debug_col1:
                        st.markdown("**Skills detected in resume:**")
                        if resume_skills:
                            st.markdown(", ".join(f"`{s}`" for s in sorted(resume_skills)))
                        else:
                            st.caption("None detected.")
                    with debug_col2:
                        st.markdown("**Skills detected in JD:**")
                        if jd_skills:
                            st.markdown(", ".join(f"`{s}`" for s in sorted(jd_skills)))
                        else:
                            st.caption("None detected.")

                    st.markdown("---")
                    st.markdown("**Extracted resume text (first 3000 chars):**")
                    st.text(resume_text[:3000] + ("..." if len(resume_text) > 3000 else ""))

# Footer
st.divider()
st.markdown("""
<div style="text-align: center; padding: 1rem 0; color: rgba(255,255,255,0.6); font-size: 0.9rem;">
    Built with ❤️ using Streamlit · pdfplumber · spaCy · scikit-learn<br>
    <span style="font-size: 0.85rem;">
        Open source · View on <a href="https://github.com/" style="color: #4fc3f7;">GitHub</a>
    </span>
</div>
""", unsafe_allow_html=True)